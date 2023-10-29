from pathlib import Path
from tempfile import TemporaryDirectory

import pytest
from docx import Document  # type: ignore

from src.office_tools.o_tool import get_installed_combinations


@pytest.fixture
def temp_doc():
    with TemporaryDirectory() as temp_dir:
        temp_file_l = Path(temp_dir) / "test.docx"
        document = Document()
        document.add_heading("Document Title", 0)
        p = document.add_paragraph("Test Text")
        p.add_run("bold run text").bold = True
        document.save(temp_file_l)

        yield temp_file_l


@pytest.mark.parametrize("office_tool_instance", get_installed_combinations())
def test_installed_tools(office_tool_instance, temp_doc):
    print(office_tool_instance)
    doc_handler = office_tool_instance.doc
    pdf_file = doc_handler.to_pdf(temp_doc)
    assert pdf_file.exists()
    ...


def test_smth():
    assert 1 > 0
